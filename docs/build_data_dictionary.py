"""
Title: Build PDA Data Dictionary (Word)
Description: Generate docs/PDA_data_dictionary.docx from live sources — the field
    definitions are written here, the dataset summary and human-review breakdown
    are computed from data/pda.db, and the extraction prompt/config are imported
    verbatim from pda.extract. Re-run after a corpus change to refresh the doc:
        .venv/bin/python docs/build_data_dictionary.py
    The human-review "reasons" are a keyword categorization of the free-text
    review_note field; a note may match more than one theme, so the theme counts
    can sum to more than the number of flagged reports.
Changelog:
    2026-06-06  Initial version (data dictionary + dataset summary + prompt).
    2026-06-06  Add human-review counts and review_note reason breakdown.
"""

import datetime
import os
import re
import sqlite3
import sys

# Make the project root importable and path-anchor outputs, so this script runs
# the same whether invoked as `docs/build_data_dictionary.py` or from elsewhere.
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if ROOT not in sys.path:
    sys.path.insert(0, ROOT)

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt

from pda.extract import MODEL, SYSTEM_PROMPT, TOOL_NAME, USER_INSTRUCTION

DB_PATH = os.path.join(ROOT, "data", "pda.db")
OUT_PATH = os.path.join(ROOT, "docs", "PDA_data_dictionary.docx")
FIELD_WIDTHS = (Inches(2.1), Inches(1.7), Inches(3.4))

# Keyword categorization of review_note. Patterns are intentionally specific
# (e.g. literal "N/A" rather than "na", which would match state names).
REVIEW_REASON_PATTERNS = {
    "Requested counties not named in the document":
        r"not named|without naming|\bunnamed\b|no county names|not individually named|no .{0,15}names? (are|were) given",
    "Missing-value markers (UNK / N/A / not conducted / no PDA)":
        r"\bUNK\b|N/A|missing-value marker|not conducted|no (joint )?PDA|left null|\bdashes?\b|illegible",
    "PA categories not specified":
        r"PA categor",
    "Legacy / nonstandard IA fields (poverty-elderly, combined SSI/SNAP)":
        r"\blegacy\b|elderly households|low income households|SSI/SNAP|combined value",
    "Narrative vs per-capita county-count discrepancy":
        r"count discrepancy|\bdiscrepancy\b|per-capita list (names|contains|plus)|but only \d+ .{0,30}named|more counties (named|listed)",
    "Source typo / malformed or anomalous value":
        r"\btypo\b|misspell|spelled|likely .{0,20}County|not a (known )?.{0,20}county|data anomaly|malformed|extra slash|drafting error",
    "Date / incident-period ambiguity (multiple or inconsistent dates)":
        r"incident period|two request dates|two separate incidents|splits damages|typo for \d{4}|inconsistent with|\bamended\b",
    "PA marked 'not requested' but per-capita indicators printed":
        r"not requested.{0,60}(indicator|per[- ]capita)|indicators? .{0,20}printed|still prints",
    "Multi-phase request, multiple governors, or non-Governor requestor":
        r"two different governors|Vice-Chair|two-phase|appeal filed by|former Gov|Tribal request|two governors",
}

REPORT_FIELDS = [
    ("GROUP", "Identity & outcome"),
    ("source_pdf", "string (PK)", "Path to the source PDF; primary key and join key for counties. (pipeline-added)"),
    ("report_outcome", "enum: Declared / Denied / Denial of Appeal, or null", "Outcome from the report's outcome line."),
    ("decision_date", "date YYYY-MM-DD, or null", "Date on the outcome line."),
    ("jurisdiction_name", "string or null", "State or tribe name."),
    ("state_abbr", "string or null", "Two-letter USPS code; null for tribes."),
    ("requestor_type", "string or null", "Governor / Tribal Chairman / Acting Chairman / etc."),
    ("requestor_name", "string or null", "Named person who requested."),
    ("incident_name", "string or null", "Incident description (e.g., \"Gwen Fire\")."),
    ("incident_begin", "date or null", "Incident start."),
    ("incident_end", "date or null", "Incident end."),
    ("request_date", "date or null", "Date the declaration was requested."),
    ("disaster_number", "integer or null", "Numeric part of e.g. FEMA-4807-DR -> 4807; null for denials/appeals."),
    ("declaration_type", "enum: DR / EM, or null", "Major disaster vs emergency; null when no number."),
    ("denial_reason", "string or null", "Stated reason (denials/appeals)."),
    ("original_denial_date", "date or null", "First denial date (appeals only)."),
    ("appeal_date", "date or null", "Appeal date (appeals only)."),
    ("GROUP", "Requested programs"),
    ("ia_requested", "boolean", "Individual Assistance requested."),
    ("pa_requested", "boolean", "Public Assistance requested."),
    ("hm_requested", "boolean", "Hazard Mitigation requested."),
    ("pa_categories_requested", "string or null", "PA categories, e.g. \"A,B\" or \"A-F\"."),
    ("GROUP", "Individual Assistance (state-level)"),
    ("ia_residences_total", "number or null", "Total residences impacted."),
    ("ia_destroyed", "number or null", "Residences destroyed."),
    ("ia_major", "number or null", "Major damage count."),
    ("ia_minor", "number or null", "Minor damage count."),
    ("ia_affected", "number or null", "Affected count."),
    ("ia_pct_insured", "number or null", "% insured residences."),
    ("ia_pct_flood_insured", "number or null", "% flood-insured (when split out)."),
    ("ia_pct_poverty", "number or null", "% poverty households (modern)."),
    ("ia_pct_ssi", "number or null", "% receiving SSI (modern)."),
    ("ia_pct_snap", "number or null", "% receiving SNAP (modern)."),
    ("ia_pct_ownership", "number or null", "% owner-occupied (modern)."),
    ("ia_unemployment", "number or null", "Unemployment rate (modern)."),
    ("ia_pct_age_65_plus", "number or null", "% age 65+ (modern)."),
    ("ia_pct_age_18_under", "number or null", "% age 18 and under (modern)."),
    ("ia_pct_disability", "number or null", "% with disability (modern)."),
    ("ia_icc_ratio", "number or null", "IHP Cost-to-Capacity ratio (modern)."),
    ("ia_pct_low_income", "number or null", "% low-income households (legacy reports)."),
    ("ia_pct_elderly", "number or null", "% elderly households (legacy reports)."),
    ("ia_cost_estimate", "number or null", "Total IA cost estimate ($)."),
    ("GROUP", "Public Assistance (state-level)"),
    ("pa_primary_impact", "string or null", "Primary PA impact description."),
    ("pa_cost_estimate", "number or null", "Total PA cost estimate ($)."),
    ("pa_statewide_per_capita", "number or null", "Statewide per-capita impact ($)."),
    ("pa_statewide_per_capita_indicator", "number or null", "Statewide per-capita threshold ($)."),
    ("pa_countywide_per_capita_indicator", "number or null", "Countywide per-capita threshold ($)."),
    ("GROUP", "Review"),
    ("needs_review", "boolean", "Model flagged the report as ambiguous/atypical."),
    ("review_note", "string or null", "Short reason for the flag."),
    ("GROUP", "Provenance / run metadata (pipeline-added)"),
    ("report_type", "string", "FEMA folder/category: MajorDisaster / Expedited / Denials / AppealDenials / Other."),
    ("url", "string or null", "Source URL from the manifest."),
    ("posted_date", "date or null", "Date posted on the FEMA index."),
    ("parser_model", "string", "Model used (e.g., claude-opus-4-8)."),
    ("extracted_at", "string (ISO 8601)", "UTC timestamp of extraction."),
]

COUNTY_FIELDS = [
    ("county_id", "integer (PK, autoincrement)", "Surrogate primary key. (pipeline-added)"),
    ("source_pdf", "string (FK -> reports.source_pdf, cascade)", "Links the unit to its report. (pipeline-added)"),
    ("county_name", "string or null", "Raw name as printed; null only when no unit is named anywhere."),
    ("geo_type", "enum: county / parish / borough / tribe / reservation / city-county / municipality / unknown", "Kind of unit; lets you filter non-county jurisdictions from the FIPS/political join."),
    ("per_capita_impact", "number or null", "Dollar figure from the PA countywide per-capita list."),
    ("requested_ia", "boolean", "Unit requested for Individual Assistance."),
    ("requested_pa", "boolean", "Unit requested for Public Assistance."),
    ("granted_ia", "boolean", "IA actually made available here (always false for denials/appeals)."),
    ("granted_pa", "boolean", "PA actually made available here (always false for denials/appeals)."),
    ("source", "enum: per_capita / narrative / both / none", "Where in the document the unit was found."),
]


def summarize_db(db_path):
    """Return dataset counts and the review-reason breakdown from the SQLite DB.

    Args:
        db_path: Path to the pda.db SQLite store.
    Returns:
        A dict with row counts, outcome counts, the flagged count, and a sorted
        list of (reason_label, count) for review_note themes plus an
        uncategorized count.
    """
    conn = sqlite3.connect(db_path)
    n_reports = conn.execute("SELECT COUNT(*) FROM reports").fetchone()[0]
    n_counties = conn.execute("SELECT COUNT(*) FROM report_counties").fetchone()[0]
    outcomes = dict(conn.execute(
        "SELECT report_outcome, COUNT(*) FROM reports GROUP BY report_outcome"))
    flagged_notes = [
        row[0] or "" for row in conn.execute(
            "SELECT review_note FROM reports WHERE needs_review = 1")
    ]
    conn.close()

    compiled = {label: re.compile(pat, re.I)
                for label, pat in REVIEW_REASON_PATTERNS.items()}
    reason_counts = sorted(
        ((label, sum(1 for n in flagged_notes if rx.search(n)))
         for label, rx in compiled.items()),
        key=lambda item: item[1], reverse=True)
    uncategorized = sum(
        1 for n in flagged_notes
        if not any(rx.search(n) for rx in compiled.values()))

    return {
        "n_reports": n_reports,
        "n_counties": n_counties,
        "declared": outcomes.get("Declared", 0),
        "denied": outcomes.get("Denied", 0),
        "appeal": outcomes.get("Denial of Appeal", 0),
        "flagged": len(flagged_notes),
        "reason_counts": reason_counts,
        "uncategorized": uncategorized,
    }


def add_field_table(doc, rows):
    """Add a 3-column Field/Type/Description table with bold group separators."""
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for cell, text in zip(table.rows[0].cells, ("Field", "Type", "Description")):
        cell.text = ""
        cell.paragraphs[0].add_run(text).bold = True
    for row in rows:
        if row[0] == "GROUP":
            merged = table.add_row().cells
            head = merged[0].merge(merged[1]).merge(merged[2])
            head.text = ""
            run = head.paragraphs[0].add_run(row[1])
            run.bold = True
            run.italic = True
        else:
            for cell, text in zip(table.add_row().cells, row):
                cell.text = ""
                cell.paragraphs[0].add_run(text)
    for row in table.rows:
        for cell, width in zip(row.cells, FIELD_WIDTHS):
            cell.width = width


def add_kv_table(doc, pairs, headers):
    """Add a simple two-column key/value table."""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = ""
        cell.paragraphs[0].add_run(text).bold = True
    for key, value in pairs:
        cells = table.add_row().cells
        cells[0].text = ""
        cells[0].paragraphs[0].add_run(str(key))
        cells[1].text = ""
        cells[1].paragraphs[0].add_run(str(value))


def add_mono_block(doc, text):
    """Render multi-line text verbatim in a monospace block (one para per line)."""
    for line in text.split("\n"):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        run = para.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)


def build(db_path=DB_PATH, out_path=OUT_PATH):
    """Build the Word data dictionary and write it to out_path."""
    s = summarize_db(db_path)
    doc = Document()

    doc.add_heading("FEMA PDA Extraction — Data Dictionary", level=0)
    intro = doc.add_paragraph()
    intro.add_run("Normalized two-table dataset (")
    intro.add_run("data/pda.db").font.name = "Courier New"
    intro.add_run(", SQLite), generated from pda/schema.py. The two tables join on ")
    intro.add_run("source_pdf").font.name = "Courier New"
    intro.add_run(".")

    # Dataset summary
    doc.add_heading("Dataset Summary", level=1)
    doc.add_paragraph(
        f"Live counts from data/pda.db as of {datetime.date.today().isoformat()}.")
    doc.add_heading("Row counts", level=2)
    add_kv_table(doc, [("reports", f"{s['n_reports']:,}"),
                       ("report_counties", f"{s['n_counties']:,}")],
                 ("Table", "Rows"))
    doc.add_heading("Reports by outcome", level=2)
    add_kv_table(doc, [
        ("Approvals (Declared)", f"{s['declared']:,}"),
        ("Denials (Denied)", f"{s['denied']:,}"),
        ("Appeal denials (Denial of Appeal)", f"{s['appeal']:,}"),
        ("All denials & appeals (treatment)", f"{s['denied'] + s['appeal']:,}"),
        ("Total reports", f"{s['n_reports']:,}"),
    ], ("Category", "Count"))

    # Human review queue
    doc.add_heading("Human review queue", level=2)
    pct = 100 * s["flagged"] / s["n_reports"] if s["n_reports"] else 0
    doc.add_paragraph(
        f"{s['flagged']:,} of {s['n_reports']:,} reports ({pct:.1f}%) are flagged "
        f"needs_review = true. The themes below categorize the free-text "
        f"review_note; a report may match more than one theme, so counts can sum "
        f"to more than the flagged total.")
    reason_rows = [(label, f"{count:,}  ({100 * count / s['flagged']:.1f}%)")
                   for label, count in s["reason_counts"]]
    reason_rows.append(
        ("Other / uncategorized note",
         f"{s['uncategorized']:,}  ({100 * s['uncategorized'] / s['flagged']:.1f}%)"))
    add_kv_table(doc, reason_rows, ("Review reason (theme)", "Reports"))

    # Field tables
    doc.add_heading("reports — one row per PDA report", level=1)
    add_field_table(doc, REPORT_FIELDS)
    doc.add_heading("report_counties — one row per (report x geographic unit)", level=1)
    add_field_table(doc, COUNTY_FIELDS)

    doc.add_heading("Notes", level=1)
    for note in [
        "The unit's state is not stored on report_counties by design — join to reports.state_abbr via source_pdf (a PDA report is single-state).",
        "In SQLite, columns are declared without a type so values keep their native storage class: booleans are 0/1, dates/enums are TEXT, numbers are REAL/INTEGER, absent values are NULL.",
        "Dates are stored as YYYY-MM-DD strings (lexically sortable); cast in your analysis tool as needed.",
    ]:
        doc.add_paragraph(note, style="List Bullet")

    # Extraction prompt & config
    doc.add_page_break()
    doc.add_heading("Extraction Prompt & Configuration", level=1)
    doc.add_paragraph(
        "The exact request used to extract each PDF. Reproduced verbatim from "
        "pda/extract.py (the single source of truth).")
    doc.add_heading("Request configuration", level=2)
    add_kv_table(doc, [
        ("Model", MODEL),
        ("Thinking", "adaptive"),
        ("Effort", "high"),
        ("max_tokens", "16000 (32000 used on retry for one oversized 49-county report)"),
        ("Output mechanism",
         f"Non-strict tool use: the model calls the {TOOL_NAME} tool (input_schema "
         f"from pda/schema.py); input is re-validated with Pydantic. Used instead "
         f"of structured outputs (16-union cap vs ~41 nullable fields)."),
        ("Input",
         "Native PDF as a base64 document block; system prompt sent as a cached "
         "(ephemeral) stable prefix."),
    ], ("Setting", "Value"))
    doc.add_heading("User instruction (user turn)", level=2)
    add_mono_block(doc, USER_INSTRUCTION)
    doc.add_heading("System prompt", level=2)
    add_mono_block(doc, SYSTEM_PROMPT)

    doc.save(out_path)
    return s


if __name__ == "__main__":
    summary = build()
    print(f"Wrote {OUT_PATH}")
    print(f"  reports={summary['n_reports']} counties={summary['n_counties']} "
          f"flagged={summary['flagged']}")
    for label, count in summary["reason_counts"]:
        print(f"  {count:4d}  {label}")
    print(f"  {summary['uncategorized']:4d}  Other / uncategorized")
