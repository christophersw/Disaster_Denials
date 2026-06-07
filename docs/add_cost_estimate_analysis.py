"""
Title: Cost-Estimate Coverage & Variability Analysis
Description: Queries data/pda.db for the IA and PA cost-estimate fields, summarizes
    how many reports carry both estimates, only one, or neither (split by outcome),
    computes the sample standard deviation of each estimate by outcome, and appends
    the results as two tables to docs/PDA_data_dictionary.docx.
Changelog:
    2026-06-07  Initial version.
    2026-06-07  Add min/max columns to the std-dev table; make the append
                idempotent so re-runs replace the section instead of duplicating it.
"""

import sqlite3
import statistics
from pathlib import Path

import docx

DB_PATH = Path("data/pda.db")
DOC_PATH = Path("docs/PDA_data_dictionary.docx")
TABLE_STYLE = "Light Grid Accent 1"
SECTION_HEADING = "Cost-Estimate Coverage and Variability"

# "Declared" reports are approvals; both denial flavors are folded into "denied".
OUTCOME_GROUPS = [
    ("Approved", ("Declared",)),
    ("Denied", ("Denied", "Denial of Appeal")),
]


def fetch_group_rows(connection, outcome_values):
    """Return (ia_cost_estimate, pa_cost_estimate) tuples for one outcome group.

    Args:
        connection: open sqlite3 connection.
        outcome_values: tuple of report_outcome strings that belong to the group.

    Returns:
        list of (ia, pa) tuples where each element is a float or None.
    """
    placeholders = ",".join("?" for _ in outcome_values)
    query = (
        "SELECT ia_cost_estimate, pa_cost_estimate "
        "FROM reports WHERE report_outcome IN (%s)" % placeholders
    )
    return connection.execute(query, outcome_values).fetchall()


def summarize_presence(rows):
    """Count how many rows have both estimates, exactly one, or neither.

    Args:
        rows: list of (ia, pa) tuples.

    Returns:
        dict with keys 'both', 'only_one', 'neither', 'total'.
    """
    both = only_one = neither = 0
    for ia, pa in rows:
        has_ia = ia is not None
        has_pa = pa is not None
        if has_ia and has_pa:
            both += 1
        elif has_ia or has_pa:
            only_one += 1
        else:
            neither += 1
    return {"both": both, "only_one": only_one, "neither": neither, "total": len(rows)}


def sample_std(values):
    """Sample standard deviation (n-1) of a list of numbers.

    Args:
        values: list of floats.

    Returns:
        float standard deviation, or None when fewer than two values are present.
    """
    return statistics.stdev(values) if len(values) > 1 else None


def money(value):
    """Format a number as a whole-dollar string, or an em dash when None."""
    return "—" if value is None else "${:,.0f}".format(value)


def add_presence_table(document, group_summaries):
    """Append the cost-estimate presence table to the document.

    Args:
        document: python-docx Document.
        group_summaries: list of (group_label, summary_dict) pairs.
    """
    headers = [
        "Outcome",
        "Both estimates",
        "Only one estimate",
        "Neither estimate",
        "Total reports",
    ]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = TABLE_STYLE
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for label, summary in group_summaries:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = str(summary["both"])
        cells[2].text = str(summary["only_one"])
        cells[3].text = str(summary["neither"])
        cells[4].text = str(summary["total"])


def add_stddev_table(document, group_stats):
    """Append the standard-deviation table to the document.

    Args:
        document: python-docx Document.
        group_stats: list of (group_label, stats_dict) pairs where stats_dict
            holds ia_n, ia_min, ia_max, ia_std and the pa_* equivalents.
    """
    headers = [
        "Outcome",
        "IA n",
        "IA min",
        "IA max",
        "IA std dev",
        "PA n",
        "PA min",
        "PA max",
        "PA std dev",
    ]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = TABLE_STYLE
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for label, stats in group_stats:
        cells = table.add_row().cells
        values = [
            label,
            str(stats["ia_n"]),
            money(stats["ia_min"]),
            money(stats["ia_max"]),
            money(stats["ia_std"]),
            str(stats["pa_n"]),
            money(stats["pa_min"]),
            money(stats["pa_max"]),
            money(stats["pa_std"]),
        ]
        for cell, text in zip(cells, values):
            cell.text = text


def remove_existing_section(document, heading_text):
    """Remove a previously-appended section so re-runs don't duplicate it.

    Deletes the heading paragraph matching heading_text and every body element
    (paragraphs and tables) that follows it. Safe because the section is always
    appended at the end of the document.

    Args:
        document: python-docx Document.
        heading_text: exact text of the section's Heading 1 paragraph.
    """
    body = document.element.body
    children = list(body)
    start_index = None
    for index, element in enumerate(children):
        if element.tag.endswith("}p") and element.text is not None:
            # Gather all text within the paragraph element.
            text = "".join(node.text or "" for node in element.iter() if node.tag.endswith("}t"))
            if text.strip() == heading_text:
                start_index = index
                break
    if start_index is None:
        return
    for element in children[start_index:]:
        # Preserve the body's trailing sectPr (section properties); removing it
        # corrupts the document's layout and breaks later add_table calls.
        if element.tag.endswith("}sectPr"):
            continue
        body.remove(element)


def main():
    connection = sqlite3.connect(DB_PATH)
    presence_summaries = []
    stddev_stats = []
    for label, outcome_values in OUTCOME_GROUPS:
        rows = fetch_group_rows(connection, outcome_values)
        presence_summaries.append((label, summarize_presence(rows)))

        ia_values = [ia for ia, _ in rows if ia is not None]
        pa_values = [pa for _, pa in rows if pa is not None]
        stddev_stats.append(
            (
                label,
                {
                    "ia_n": len(ia_values),
                    "ia_min": min(ia_values) if ia_values else None,
                    "ia_max": max(ia_values) if ia_values else None,
                    "ia_std": sample_std(ia_values),
                    "pa_n": len(pa_values),
                    "pa_min": min(pa_values) if pa_values else None,
                    "pa_max": max(pa_values) if pa_values else None,
                    "pa_std": sample_std(pa_values),
                },
            )
        )
    connection.close()

    document = docx.Document(DOC_PATH)
    remove_existing_section(document, SECTION_HEADING)
    document.add_heading(SECTION_HEADING, level=1)
    document.add_paragraph(
        "Analysis of the ia_cost_estimate and pa_cost_estimate fields across all "
        "1,378 reports. “Approved” covers reports with outcome “Declared”; "
        "“Denied” combines “Denied” and “Denial of Appeal”. UNK/Unknown "
        "values were already stored as null and are counted as missing."
    )

    document.add_heading("Estimate coverage by outcome", level=2)
    document.add_paragraph(
        "Counts of reports that report both cost estimates, exactly one, or neither."
    )
    add_presence_table(document, presence_summaries)

    document.add_heading("Spread of cost estimates by outcome", level=2)
    document.add_paragraph(
        "Minimum, maximum, and sample standard deviation (n−1) of each cost "
        "estimate in U.S. dollars, computed over the non-null values only; n is "
        "the count of reports contributing to each figure."
    )
    add_stddev_table(document, stddev_stats)

    document.save(DOC_PATH)

    # Echo results to the console for verification.
    print("Presence by outcome:")
    for label, summary in presence_summaries:
        print(f"  {label}: {summary}")
    print("Spread by outcome:")
    for label, stats in stddev_stats:
        print(
            f"  {label}: IA n={stats['ia_n']} min={money(stats['ia_min'])} "
            f"max={money(stats['ia_max'])} std={money(stats['ia_std'])} | "
            f"PA n={stats['pa_n']} min={money(stats['pa_min'])} "
            f"max={money(stats['pa_max'])} std={money(stats['pa_std'])}"
        )


if __name__ == "__main__":
    main()
